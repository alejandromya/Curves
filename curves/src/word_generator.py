from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from src.tables_generator import generar_tablas_combinadas

def generar_word_unico(bloques, output_doc="INFORME_FINAL.docx"):
    doc = Document()

    # Función auxiliar para añadir párrafos con fuente y tamaño personalizado
    def add_paragraph(text, level=None, bold=False, font_name='Arial', font_size=10, align=None):
        if level:
            p = doc.add_heading(text, level=level)
        else:
            p = doc.add_paragraph(text)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        if align:
            p.alignment = align
        return p

    # Título principal
    add_paragraph("Resultados", level=1, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    doc.add_paragraph()  # espacio

    # Generar tablas combinadas
    (headers_cyclic, filas_cyclic), (headers_3rd, filas_3rd) = generar_tablas_combinadas(bloques)

    # Función auxiliar para crear tablas
    def add_table(headers, filas):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Encabezado
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header)
            hdr_run = hdr_cells[i].paragraphs[0].runs[0]
            hdr_run.bold = True
            hdr_run.font.name = 'Arial'
            hdr_run.font.size = Pt(10)
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Filas
        for fila in filas:
            row_cells = table.add_row().cells
            for i, value in enumerate(fila):
                row_cells[i].text = str(value)
                row_run = row_cells[i].paragraphs[0].runs[0]
                row_run.font.name = 'Arial'
                row_run.font.size = Pt(10)
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # espacio
        return table

    # Agregar tablas
    add_table(headers_cyclic, filas_cyclic)
    add_table(headers_3rd, filas_3rd)

    # Gráficas por bloque
    for bloque in bloques:
        doc.add_page_break()
        add_paragraph(bloque['titulo'], level=2, align=WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)
        doc.add_paragraph()  # espacio
        doc.add_picture(bloque['grafico'], width=Inches(6))  # ajustar ancho de imagen

    # Guardar documento
    doc.save(output_doc)
    return output_doc
